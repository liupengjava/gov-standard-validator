#!/usr/bin/env python3
"""Build a Word validation report from the structured plain-text report."""
import argparse
import json
import os
import re

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def setup_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.4)
    sec.right_margin = Cm(2.4)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)

    for style_name in ["Heading 1", "Heading 2"]:
      style = doc.styles[style_name]
      style.font.name = "宋体"
      style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
      style.font.color.rgb = RGBColor(0, 0, 0)


def add_title(doc: Document, text: str, size: int) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = True
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)


def add_para(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Cm(0.74)
    para.paragraph_format.line_spacing = 1.25
    para.paragraph_format.space_after = Pt(4)
    if "：" in text and len(text.split("：", 1)[0]) <= 12:
        prefix, rest = text.split("：", 1)
        run = para.add_run(prefix + "：")
        run.bold = True
        para.add_run(rest)
    else:
        para.add_run(text)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    para = doc.add_heading(text, level=level)
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(6)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_index, row in enumerate(rows):
        for c_index in range(col_count):
            cell = table.cell(r_index, c_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = row[c_index] if c_index < len(row) else ""
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    run.font.name = "宋体"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    run.font.size = Pt(9)
                    if r_index == 0:
                        run.bold = True
            if r_index == 0:
                set_cell_shading(cell, "D9EAF7")


def split_table_line(line: str) -> list[str]:
    return [part.strip() for part in line.split("｜")]


def flush_table(doc: Document, table_lines: list[str]) -> None:
    if table_lines:
        add_table(doc, [split_table_line(line) for line in table_lines])
        doc.add_paragraph()
        table_lines.clear()


def build_docx(report_text: str, output_path: str) -> None:
    doc = Document()
    setup_doc(doc)
    table_lines: list[str] = []
    title_seen = False

    for raw_line in report_text.splitlines():
        line = raw_line.strip()
        if not line:
            flush_table(doc, table_lines)
            continue
        if "｜" in line and (table_lines or not re.match(r"^[VR]-\d{3}｜", line)):
            table_lines.append(line)
            continue
        flush_table(doc, table_lines)

        if line == "标准验证意见报告" and not title_seen:
            add_title(doc, line, 20)
            title_seen = True
        elif line.startswith("《") and line.endswith("》"):
            add_title(doc, line, 14)
        elif re.match(r"^[一二三四五六七八九十]+、", line):
            add_heading(doc, line, 1)
        elif re.match(r"^[VR]-\d{3}｜", line):
            add_heading(doc, line, 2)
        else:
            add_para(doc, line)

    flush_table(doc, table_lines)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--in", dest="input", required=True)
    parser.add_argument("--out", dest="output", required=True)
    args = parser.parse_args()

    with open(args.input, "r", encoding="utf-8") as f:
        payload = json.load(f)
    report_text = str(payload.get("reportText") or "").strip()
    if not report_text:
        raise SystemExit("reportText is required")
    os.makedirs(os.path.dirname(args.output), exist_ok=True)
    build_docx(report_text, args.output)
    print(args.output)


if __name__ == "__main__":
    main()
